"""

Take a list of names and make a docx custom invite for each one

"""

import docx

with open('guests.txt') as f:
    names = f.readlines()
    document = docx.Document()

    for name in names:
        name = name.strip()

        document.add_paragraph('It would be a pleasure to have the company of', style='Normal')
        document.add_paragraph(name, style='Normal')
        document.add_paragraph('at Bhayadanga bazar', style='Normal')
        document.add_paragraph('June 1st', style='Normal')
        document.add_paragraph("at 9 o'clock", style='Normal')

        document.add_page_break()

    document.save('invites.docx')
    print('Files has been created and saved as invites.docx')